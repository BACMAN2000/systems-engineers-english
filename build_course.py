# -*- coding: utf-8 -*-
"""
Systems Engineers English Course A1-C2 - Malla Curricular
Autor: Paolo Baca (NIS) - asistido por Claude
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTFILE = r"E:\CLAUDE\PAOLO BACA MANRIQUE\Systems_Engineers_English_Course_A1_C2.docx"

doc = Document()

# ---------- estilos base ----------
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

for h, sz in (('Heading 1', 20), ('Heading 2', 15), ('Heading 3', 13)):
    s = doc.styles[h]
    s.font.name = 'Arial'
    s.font.size = Pt(sz)
    s.font.color.rgb = RGBColor(0x0B, 0x3C, 0x5D)
    s.font.bold = True

# márgenes
for section in doc.sections:
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, italic=False, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if align: p.alignment = align
    return p

def add_table(headers, rows, col_widths=None, header_fill="0B3C5D", header_text_white=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        if header_text_white:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        shade_cell(hdr[i], header_fill)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.name = 'Arial'
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return t

def page_break():
    doc.add_page_break()

# =========================================================
# PORTADA
# =========================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\nSYSTEMS ENGINEERS\nENGLISH COURSE")
r.font.size = Pt(32); r.bold = True
r.font.color.rgb = RGBColor(0x0B, 0x3C, 0x5D)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A1 \u2192 C2  |  CEFR-aligned")
r.font.size = Pt(18); r.bold = True

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Malla curricular \u2013 120 clases\nESP / Cybersecurity / Emerging Technology")
r.font.size = Pt(14); r.italic = True

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\nCambridge KET \u2022 PET \u2022 FCE \u2022 CAE \u2022 CPE\nIELTS \u2022 TOEFL iBT")
r.font.size = Pt(12); r.bold = True

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\nAuthor: Paolo Baca Manrique\nNordic International School, Per\u00fa \u2013 2026")
r.font.size = Pt(11)
page_break()

# =========================================================
# 1. INTRODUCCI\u00d3N Y ENFOQUE
# =========================================================
add_heading("1. Introducci\u00f3n y enfoque pedag\u00f3gico", 1)
add_para(
    "Este programa ha sido dise\u00f1ado para ingenieros de sistemas (en formaci\u00f3n o en ejercicio) "
    "siguiendo estrictamente el Marco Com\u00fan Europeo de Referencia para las Lenguas (MCER / CEFR) "
    "del Consejo de Europa, y calibrado contra los descriptores de los ex\u00e1menes internacionales "
    "de Cambridge Assessment English, IELTS Academic y TOEFL iBT. Cada uno de los seis niveles "
    "(A1, A2, B1, B2, C1, C2) contiene 20 clases de 90 minutos, 1200 minutos por nivel, "
    "para un total de 120 clases y 180 horas did\u00e1cticas."
)
add_para(
    "El enfoque es CLIL / ESP (Content and Language Integrated Learning / English for Specific "
    "Purposes) centrado en tres l\u00edneas tem\u00e1ticas: (i) fundamentos de ingenier\u00eda de sistemas, "
    "(ii) ciberseguridad, y (iii) tecnolog\u00edas emergentes (IA/ML, cloud, IoT, blockchain, "
    "computaci\u00f3n cu\u00e1ntica). Todas las lecturas provienen o se modelan sobre fuentes de \u00faltima "
    "tecnolog\u00eda: IEEE Spectrum, ACM, NIST, CISA, arXiv, MIT Technology Review, The Register."
)
add_para(
    "Cada clase integra las cuatro destrezas (reading, listening, speaking, writing) m\u00e1s "
    "vocabulario y gram\u00e1tica en uso, y cierra con una tarea de transferencia profesional "
    "(use-of-English aplicado al mundo IT)."
)

page_break()

# =========================================================
# 2. EQUIVALENCIAS Y CUOTAS CUANTITATIVAS
# =========================================================
add_heading("2. Equivalencias CEFR \u2013 Cambridge \u2013 IELTS \u2013 TOEFL", 1)
add_para(
    "Las equivalencias siguen los mapeos oficiales de Cambridge Assessment English, "
    "IELTS Partnership (British Council / IDP / Cambridge) y ETS (TOEFL). "
    "Los rangos TOEFL/IELTS son indicativos: los ex\u00e1menes oficiales son multinivel."
)

add_table(
    ["CEFR", "Cambridge", "IELTS", "TOEFL iBT", "Perfil profesional IT"],
    [
        ["A1", "Pre-A1 Starters / A1 Movers / KET (parte baja)", "<4.0", "0\u201331", "Usuario b\u00e1sico; sigue men\u00fas y mensajes cortos"],
        ["A2", "A2 Key (KET) / A2 Flyers", "4.0", "32\u201341", "T\u00e9cnico junior; describe tareas simples de soporte"],
        ["B1", "B1 Preliminary (PET)", "4.5 \u2013 5.0", "42\u201371", "T\u00e9cnico; redacta tickets, lee documentaci\u00f3n guiada"],
        ["B2", "B2 First (FCE)", "5.5 \u2013 6.5", "72\u201394", "Ingeniero; reuniones t\u00e9cnicas, informes, mitigaci\u00f3n"],
        ["C1", "C1 Advanced (CAE)", "7.0 \u2013 8.0", "95\u2013120", "Ingeniero senior / arquitecto; whitepapers, lid. t\u00e9cnico"],
        ["C2", "C2 Proficiency (CPE)", "8.5 \u2013 9.0", "118\u2013120", "Investigador / CISO / autor; publicaciones, keynotes"],
    ],
    col_widths=[1.3, 4.2, 2.0, 2.0, 6.5],
)

doc.add_paragraph()
add_heading("2.1 Cuotas cuantitativas por nivel", 2)
add_para(
    "Las cifras siguen los descriptores English Profile del British Council / Cambridge y el "
    "Common European Framework: word lists (receptive + productive), Cambridge writing word "
    "count bands y rangos de velocidad lectora (wpm)."
)

add_table(
    ["Nivel", "Vocab. acumulado", "Vocab. nuevo x nivel", "Escritura (palabras/tarea)", "Lectura (palabras/texto)", "Velocidad lectora (wpm)", "Horas de listening"],
    [
        ["A1", "500", "+500", "25\u201335 (notas, formularios)", "100\u2013180", "80\u2013110", "15 h"],
        ["A2", "1 000", "+500", "35\u201350 (emails cortos)", "180\u2013300", "110\u2013140", "20 h"],
        ["B1", "2 000", "+1 000", "100\u2013140 (email/report)", "300\u2013500", "140\u2013170", "25 h"],
        ["B2", "4 000", "+2 000", "140\u2013190 (essay/report)", "550\u2013800", "170\u2013220", "30 h"],
        ["C1", "8 000", "+4 000", "220\u2013260 (essay/proposal)", "800\u20131 100", "220\u2013260", "35 h"],
        ["C2", "16 000+", "+8 000", "280\u2013350 (paper/review)", "1 100\u20131 500", "260\u2013300", "40 h"],
    ],
    col_widths=[1.3, 2.2, 2.6, 3.8, 2.8, 2.5, 1.8],
)

doc.add_paragraph()
add_heading("2.2 Modelo de 20 clases por nivel", 2)
add_para(
    "Cada nivel se organiza en 4 m\u00f3dulos de 5 clases. La clase 20 es evaluaci\u00f3n sumativa "
    "en formato Cambridge completo (Reading & Use of English + Writing + Listening + Speaking)."
)
add_table(
    ["M\u00f3dulo", "Clases", "Foco"],
    [
        ["M1 Foundations", "1\u20135", "Vocabulario nuclear + gram\u00e1tica objetivo del nivel"],
        ["M2 Core IT / ESP", "6\u201310", "Ingenier\u00eda de sistemas aplicada"],
        ["M3 Cybersecurity", "11\u201315", "Amenazas, defensa, pol\u00edticas"],
        ["M4 Emerging Tech + Exam Prep", "16\u201320", "IA, cloud, IoT, blockchain + mock exam"],
    ],
    col_widths=[4.0, 2.0, 10.0],
)
page_break()

# =========================================================
# 3. FUENTES DE AUDIO SIN COPYRIGHT
# =========================================================
add_heading("3. Fuentes de audio sin copyright / uso libre", 1)
add_para(
    "Las siguientes fuentes permiten uso educativo libre porque son (a) obras del gobierno "
    "federal de EE.\u202fUU. y por lo tanto dominio p\u00fablico, (b) dominio p\u00fablico por antig\u00fcedad "
    "o donaci\u00f3n, o (c) licenciadas bajo Creative Commons (CC BY, CC BY-SA o CC0). "
    "Verifique siempre la licencia espec\u00edfica del archivo antes de redistribuirlo."
)

add_table(
    ["Fuente", "Dominio", "Licencia", "Tipo de contenido", "Nivel CEFR \u00fatil"],
    [
        ["VOA Learning English", "learningenglish.voanews.com", "Dominio p\u00fablico (US gov)", "Noticias de tecnolog\u00eda y ciencia a 3 velocidades (Level 1/2/3)", "A1\u2013B2"],
        ["NASA Audio / Podcasts", "nasa.gov/audio", "Dominio p\u00fablico (US gov)", "Ciencia, espacio, ingenier\u00eda", "B1\u2013C2"],
        ["CISA Cyber Podcasts", "cisa.gov", "Dominio p\u00fablico (US gov)", "Alertas y charlas de ciberseguridad", "B2\u2013C2"],
        ["NIST Cybersecurity", "nist.gov", "Dominio p\u00fablico (US gov)", "Est\u00e1ndares, frameworks, webinars", "B2\u2013C2"],
        ["FBI Cyber / Gov Podcasts", "fbi.gov", "Dominio p\u00fablico (US gov)", "Casos de cibercrimen", "B2\u2013C1"],
        ["LibriVox", "librivox.org", "Dominio p\u00fablico", "Audiolibros (Verne, Wells, Asimov*)", "A2\u2013C2"],
        ["Project Gutenberg Audio", "gutenberg.org", "Dominio p\u00fablico", "Textos cl\u00e1sicos con audio", "B1\u2013C2"],
        ["Internet Archive", "archive.org", "Mixto (filtrar CC / PD)", "Conferencias, radio antigua, confs. t\u00e9cnicas", "Todos"],
        ["MIT OpenCourseWare", "ocw.mit.edu", "CC BY-NC-SA", "Clases completas de ingenier\u00eda y CS", "B2\u2013C2"],
        ["OpenLearn (Open University)", "open.edu/openlearn", "CC BY-NC-SA", "Cursos y podcasts", "B1\u2013C1"],
        ["TED Talks / TED-Ed", "ted.com", "CC BY-NC-ND (verificar)", "Charlas cortas con transcripci\u00f3n", "B1\u2013C2"],
        ["Wikimedia Commons Audio", "commons.wikimedia.org", "CC BY-SA / PD", "Art\u00edculos hablados, efectos", "A2\u2013C2"],
        ["Freesound.org", "freesound.org", "CC0 / CC BY", "Efectos, ambientes (keyboard, server hum)", "Todos (aux.)"],
        ["Pixabay Audio", "pixabay.com/music", "Pixabay Content License", "M\u00fasica y SFX sin atribuci\u00f3n", "Todos (aux.)"],
        ["Librivox + Hacker News Audio", "librivox / hn", "PD / CC", "Lecturas de ensayos tecnol\u00f3gicos", "C1\u2013C2"],
        ["Berkman Klein Center", "cyber.harvard.edu", "CC BY", "Podcasts sobre Internet y sociedad", "C1\u2013C2"],
        ["EFF Podcasts", "eff.org", "CC BY", "Privacidad, DRM, pol\u00edtica digital", "B2\u2013C2"],
    ],
    col_widths=[3.8, 3.5, 2.8, 4.2, 1.7],
)
add_para(
    "*Nota: obras de Asimov entran progresivamente al dominio p\u00fablico. Verificar jurisdicci\u00f3n "
    "antes de uso. Para Per\u00fa, rige el plazo de vida del autor + 70 a\u00f1os.",
    italic=True, size=9
)

page_break()

# =========================================================
# 4. COMPETENCIAS POR DESTREZA (MATRIZ)
# =========================================================
add_heading("4. Competencias por destreza y nivel", 1)
add_table(
    ["Nivel", "Reading (Cambridge style)", "Listening (tareas)", "Writing (tipo y palabras)", "Speaking"],
    [
        ["A1",
         "Se\u00f1ales, men\u00fas, mensajes cortos (100\u2013180 w). Gap-fill 5 items.",
         "Di\u00e1logos lentos; match image-word; 0.8x speed.",
         "Formulario + nota (25\u201335 w).",
         "Saludo, deletrear URL, describir rutina IT."],
        ["A2",
         "Emails, posts de blog sencillos (180\u2013300 w). MCQ 5\u20137.",
         "Anuncios, mensajes de voz; gap-fill.",
         "Email informativo / descripci\u00f3n (35\u201350 w).",
         "Describir dispositivo, pedir ayuda t\u00e9cnica."],
        ["B1",
         "Textos divulgativos (300\u2013500 w). Matching headings, T/F/NG.",
         "Entrevistas cortas, podcast VOA nivel 2.",
         "Email formal / reporte (100\u2013140 w).",
         "Presentar un proyecto en 2 min; negociaci\u00f3n simple."],
        ["B2",
         "Art\u00edculos de tech news (550\u2013800 w). Multiple matching, open cloze.",
         "Conferencia corta TED-Ed (6\u201310 min).",
         "Essay / review / report (140\u2013190 w).",
         "Debate estructurado 4 min, pros/cons."],
        ["C1",
         "Whitepapers, art\u00edculos IEEE Spectrum (800\u20131 100 w). Keyword transformations, summary.",
         "Keynotes de confs. (15\u201325 min), m\u00faltiples acentos.",
         "Essay / proposal / formal report (220\u2013260 w).",
         "Presentaci\u00f3n 6\u20138 min + Q&A."],
        ["C2",
         "Papers arXiv / ACM (1 100\u20131 500 w). Summary + comparative.",
         "Paneles, debate acad\u00e9mico, jerga de investigaci\u00f3n.",
         "Article / paper / critical review (280\u2013350 w).",
         "Defender tesis 10 min ante comit\u00e9."],
    ],
    col_widths=[1.3, 5.0, 3.5, 3.8, 3.4],
)

page_break()

# =========================================================
# 5. MALLA DETALLADA
# =========================================================
add_heading("5. Malla curricular detallada", 1)
add_para(
    "Cada clase dura 90 minutos. Columnas: Tema / Gram\u00e1tica (Gram) / Vocabulario objetivo (V) / "
    "Reading (R, tipo Cambridge + equivalente TOEFL/IELTS) / Listening (L) / Writing (W, con "
    "conteo exacto de palabras) / Ciber/Tech angle."
)

LEVELS = [
    ("A1", "Beginner \u2013 KET (parte baja)", "500 palabras receptivas objetivo", [
        # 1..20: (tema, gram, vocab nuevas, reading, listening, writing wordcount, ciber angle)
        ("Hello, I'm a systems engineer",
         "Verb BE, subject pronouns, a/an",
         "25 palabras: name, job, company, computer, laptop, desk, office, email, phone, hi, hello, engineer, team, manager, IT, tech, help desk, user, password, login, file, folder, screen, mouse, keyboard",
         "Perfiles cortos de empleados IT (150 w). Task: match name-job. KET Reading Part 1.",
         "VOA Learning English \u2013 Level 1 'My Job' (2 min, 0.75x).",
         "Perfil personal (30 w)",
         "Rutina del helpdesk"),
        ("The alphabet of URLs and emails",
         "Alphabet, spelling, basic plurals",
         "25: dot, slash, at, underscore, hyphen, www, com, edu, net, org, http, https, link, site, page, web, browser, open, close, click, tab, window, home, back, forward",
         "URLs corporativas y direcciones de email (100 w). Gap-fill. IELTS A1 equiv.",
         "Spelling bee t\u00e9cnico (audio gobierno US, 3 min).",
         "Deletrear 5 URLs (25 w)",
         "Spoofing de dominios: diferenciar 'paypal.com' de 'paypa1.com'"),
        ("Numbers, versions and dates",
         "Cardinals 0\u2013100, ordinals, prepositions of time",
         "25: zero, one, \u2026 hundred, first, second, version, update, release, date, day, month, year, morning, afternoon, evening, today, tomorrow, yesterday, now, later, soon, hour, minute, second",
         "Release notes simplificadas (120 w). T/F.",
         "VOA Tech Report Level 1 sobre Windows update.",
         "Describir cronograma de updates (30 w)",
         "Parches de seguridad: por qu\u00e9 'Patch Tuesday' importa"),
        ("Hardware parts",
         "There is / there are, demonstratives",
         "30: CPU, RAM, hard disk, SSD, motherboard, fan, cable, port, USB, HDMI, monitor, case, power, battery, charger, plug, socket, printer, scanner, speaker, webcam, microphone, headset, dock, server, rack, tower, slot, chip, board",
         "Descripci\u00f3n de un PC de oficina (160 w). Etiquetado de imagen.",
         "Audio de NIST 'Inside a computer' (2 min, adaptado).",
         "Listar 10 componentes (30 w)",
         "Hardware tokens (YubiKey) como intro a 2FA"),
        ("Software basics",
         "Can / can't for ability",
         "25: software, program, app, install, uninstall, update, open, save, delete, edit, copy, paste, cut, print, close, start, stop, run, click, double-click, right-click, select, menu, icon, folder",
         "Manual r\u00e1pido de una app (150 w). MCQ.",
         "Tutorial corto MIT OCW 6.0001 intro (adaptado 3 min).",
         "Instrucciones para abrir un archivo (30 w)",
         "Software pirata = malware: riesgos para el A1"),
        ("My workspace",
         "Prepositions of place, possessive adjectives",
         "20: on, in, under, next to, between, in front of, behind, desk, chair, lamp, drawer, shelf, wall, window, door, my, your, his, her, our, their",
         "Descripci\u00f3n de un NOC (140 w). Matching.",
         "Ambient sounds de data center (Freesound CC0) + narraci\u00f3n.",
         "Describir tu escritorio (30 w)",
         "Clean desk policy b\u00e1sico"),
        ("Daily routines in IT",
         "Present simple, adverbs of frequency",
         "25: check, read, reply, send, join, attend, monitor, report, ticket, call, write, test, deploy, review, always, usually, often, sometimes, rarely, never, every day, daily, morning shift, night shift, break",
         "Horario de un SysAdmin junior (180 w). Gap-fill.",
         "Entrevista 'A day in IT' (OpenLearn adaptado).",
         "Mi rutina IT (35 w)",
         "Shift handover y por qu\u00e9 importa a la seguridad"),
        ("Files and folders",
         "Have / has got, some / any",
         "25: file, folder, document, image, photo, video, audio, pdf, txt, doc, xls, zip, size, name, type, path, drive, local, cloud, shared, private, public, read, write, execute",
         "Estructura de carpetas (150 w). Match paths.",
         "LibriVox \u2013 cap\u00edtulo corto sobre archivado.",
         "Organizar 10 archivos (30 w)",
         "File permissions: read/write/execute como intro UNIX"),
        ("Internet basics",
         "Present continuous",
         "25: internet, web, site, page, link, search, browser, Wi-Fi, router, modem, signal, connect, disconnect, online, offline, upload, download, stream, load, loading, speed, slow, fast, mobile data, data",
         "Gu\u00eda de conexi\u00f3n Wi-Fi (160 w). T/F.",
         "VOA Level 1 \u2013 'How the Internet works' (3 min).",
         "Explicar 'how I connect' (35 w)",
         "Wi-Fi p\u00fablico: riesgos para un usuario b\u00e1sico"),
        ("Writing safe passwords",
         "Imperatives, should / shouldn't",
         "25: password, strong, weak, long, short, letter, number, symbol, capital, lower, uppercase, change, remember, forget, reset, manager, keep, share, safe, secure, risk, danger, protect, login, logout",
         "Pol\u00edtica de contrase\u00f1as para novatos (170 w). Gap-fill.",
         "Podcast CISA 'Make a strong password' (adaptado 3 min).",
         "Escribir 5 reglas (35 w)",
         "Password strength meters; passphrases"),
        ("Mobile devices",
         "Object pronouns",
         "25: phone, smartphone, tablet, watch, app, screen, touch, swipe, tap, lock, unlock, fingerprint, face, notification, battery, charger, update, Android, iOS, settings, Bluetooth, pair, ring, silent, vibrate",
         "Gu\u00eda r\u00e1pida del iPhone (150 w). MCQ.",
         "Tutorial de accesibilidad (Apple, materiales gubernamentales US).",
         "Describir tu tel\u00e9fono (35 w)",
         "MDM y bloqueo remoto en el trabajo"),
        ("Colors and UI",
         "Adjective + noun order",
         "20: red, blue, green, yellow, black, white, gray, dark, light, button, bar, box, icon, label, text, big, small, wide, narrow, flat",
         "Gu\u00eda de dise\u00f1o b\u00e1sica (140 w). Match color-meaning.",
         "TED-Ed 'Color and UI' (adaptado).",
         "Describir una pantalla (30 w)",
         "Rojo/verde en alerts: sem\u00e1ntica visual"),
        ("Office tools",
         "Have to, need to",
         "25: word processor, spreadsheet, slide, presentation, chart, table, cell, row, column, format, font, size, bold, italic, underline, save as, export, print, copy, paste, insert, delete, undo, redo",
         "Tutorial LibreOffice (170 w). Gap-fill.",
         "Intro a Google Docs (VOA).",
         "Guiar a un colega (40 w)",
         "Macros maliciosas en Office: intro"),
        ("Email etiquette 101",
         "Politeness: could you, please",
         "25: subject, to, cc, bcc, reply, forward, attach, attachment, open, save, hi, dear, regards, best, thanks, please, sorry, send, receive, spam, junk, inbox, outbox, draft, signature",
         "Tres emails de soporte (180 w). Matching tono.",
         "Audio VOA sobre email etiquette.",
         "Email al helpdesk (40 w)",
         "Phishing visual: red flags para principiantes"),
        ("Time, meetings and schedules",
         "Prepositions of time (at, on, in)",
         "20: meeting, schedule, calendar, invite, accept, decline, reschedule, on time, late, early, busy, free, available, morning, afternoon, noon, before, after, during, between",
         "Agenda semanal de un equipo (150 w). T/F.",
         "Voicemail t\u00edpico en ingl\u00e9s.",
         "Proponer una reuni\u00f3n (40 w)",
         "Ingenier\u00eda social por tel\u00e9fono: vishing intro"),
        ("Sizes and units",
         "Countable / uncountable, much / many",
         "20: byte, kilobyte, megabyte, gigabyte, terabyte, bit, pixel, resolution, dpi, hz, mhz, ghz, storage, memory, capacity, free, used, full, empty, amount",
         "Ficha t\u00e9cnica de un laptop (170 w). Matching.",
         "VOA sobre storage.",
         "Comparar dos laptops (40 w)",
         "Por qu\u00e9 el tama\u00f1o del log importa"),
        ("Simple troubleshooting",
         "Why / because, first / then / finally",
         "25: problem, issue, error, bug, crash, freeze, restart, reboot, reinstall, check, test, fix, solve, turn on, turn off, plug in, unplug, try, work, not work, slow, fast, hot, cold, noise",
         "Di\u00e1logo helpdesk (200 w). MCQ.",
         "Audio adaptado 'Have you tried turning it off and on?' (gobierno US, dominio p\u00fablico).",
         "Pasos de troubleshooting (45 w)",
         "No dejar el equipo desbloqueado al reiniciar"),
        ("Jobs in IT",
         "Comparatives short adj.",
         "20: developer, tester, analyst, engineer, support, admin, architect, designer, devops, intern, senior, junior, lead, manager, easy, hard, big, small, busy, good",
         "Anuncios laborales (180 w). Matching job-description.",
         "Mini entrevistas (Internet Archive audio).",
         "Tu rol ideal (45 w)",
         "Junior pentester vs junior SOC analyst"),
        ("My first English CV",
         "Past simple (regular), last / ago",
         "25: studied, worked, taught, trained, learned, used, made, built, fixed, installed, tested, led, helped, year, month, university, college, degree, certificate, skill, tool, software, language, experience, reference",
         "Modelo de CV junior (180 w). Gap-fill.",
         "Podcast VOA Everyday Grammar.",
         "CV de 50 w",
         "Por qu\u00e9 no poner info sensible en el CV"),
        ("Final project + Mock KET",
         "Revisi\u00f3n integral",
         "500 acumuladas",
         "KET Reading completo (300 w total).",
         "KET Listening completo (25 min).",
         "Formulario + mensaje (35 w total)",
         "Reflexi\u00f3n: seguridad del principiante"),
    ]),

    # ================================================= A2 =================================================
    ("A2", "Elementary \u2013 KET", "1 000 palabras acumuladas", [
        ("Describing IT jobs",
         "Present simple vs continuous",
         "+50: developer, analyst, architect, responsible for, in charge of, report to, supervise, deliver, support, maintain, implement, configure, deploy, document, train, negotiate, design, code, debug, release, handover, onboard, offboard, escalate, mentor, coach, align, oversee, run, plan, lead, coordinate, decide, approve, reject, review, audit, test, verify, validate, present, summarize, discuss, propose, suggest, recommend, assign, prioritize, schedule, monitor",
         "'What does a DevOps engineer do?' (250 w). Matching.",
         "VOA Learning English \u2013 'Tech careers' Level 2.",
         "Describir un rol (45 w)",
         "Rol del SOC analyst"),
        ("My work history",
         "Past simple regular + irregular",
         "+45: began, started, finished, ended, joined, left, moved, became, built, wrote, sent, read, got, took, made, held, won, lost, found, taught, learnt, met, ran, hired, fired, quit, applied, promoted, transferred, worked on, worked with, delivered, solved, rolled out, shipped, launched, supported, trained, led, assisted, debugged, deployed, tested, documented, improved",
         "Perfil LinkedIn-style (280 w). T/F/NG.",
         "Podcast OpenLearn 'My first tech job'.",
         "3 p\u00e1rrafos experiencia (50 w)",
         "No incluir en redes datos explotables"),
        ("Operating systems compared",
         "Comparative & superlative",
         "+40: Windows, Linux, macOS, distro, kernel, shell, terminal, GUI, command line, open source, proprietary, free, paid, faster, slower, better, worse, easier, harder, most, least, popular, secure, user-friendly, flexible, powerful, lightweight, heavy, stable, unstable, reliable, support, driver, update, upgrade, install, patch, license, supported, deprecated, obsolete",
         "'Linux vs Windows server' (300 w). Multiple matching.",
         "MIT OCW clase corta sobre SO (adaptada).",
         "Ensayo comparativo (60 w)",
         "Supply chain: cada SO tiene su cadena"),
        ("Installing software",
         "Sequence: first, then, after that, finally",
         "+40: installer, setup, wizard, license, terms, accept, agree, decline, next, back, cancel, path, directory, default, custom, typical, minimum, full, component, feature, disk space, requirements, admin, rights, elevate, UAC, sudo, msi, exe, dmg, deb, rpm, repo, download, verify, hash, signature, launch, run, restart",
         "Gu\u00eda de instalaci\u00f3n (300 w). Order steps.",
         "Tutorial apt-get (adaptado).",
         "Pasos de instalaci\u00f3n (55 w)",
         "Verificar hash SHA-256 antes de instalar"),
        ("Web browsers",
         "Modals of advice: should / ought to",
         "+40: browser, tab, history, cache, cookie, bookmark, favorite, extension, plugin, add-on, incognito, private, clear, block, allow, setting, homepage, search engine, default, sync, profile, guest, address bar, bar, toolbar, download, upload, script, popup, tracker, ad, blocker, session, autofill, save, remember, sign in, sign out, account",
         "Seguridad b\u00e1sica del browser (320 w). MCQ.",
         "EFF podcast 'Browser privacy' nivel 2.",
         "Gu\u00eda (55 w)",
         "Cookies y fingerprinting"),
        ("Social media safety",
         "Must / mustn't / can't",
         "+40: profile, post, share, like, follow, friend, unfriend, block, report, privacy, public, private, settings, tag, mention, hashtag, feed, story, message, DM, group, chat, photo, video, stream, live, platform, verified, fake, impostor, scam, spam, phishing, link, hide, delete, edit, pin, archive, disable",
         "Caso: deepfake en Instagram (320 w). T/F/NG.",
         "VOA 'Be careful online' Level 2.",
         "5 consejos (60 w)",
         "OSINT sobre tus propias redes"),
        ("Networks: Wi-Fi and LAN",
         "Adverbs of manner",
         "+45: network, LAN, WAN, wireless, wired, cable, ethernet, Wi-Fi, SSID, WPA, WPA2, WPA3, open network, password, router, switch, hub, modem, access point, IP, address, public, private, range, signal, strong, weak, fast, slow, quickly, slowly, carefully, successfully, connect, disconnect, bridge, repeater, mesh, guest, admin, dashboard, login",
         "Configurando un router (320 w). Gap-fill.",
         "CISA audio sobre routers.",
         "Gu\u00eda (60 w)",
         "Firmware del router y default passwords"),
        ("Mobile devices advanced",
         "Going to (plans)",
         "+40: roaming, data plan, eSIM, carrier, operator, 4G, 5G, LTE, hotspot, tether, backup, restore, factory reset, wipe, MDM, managed, supervised, kiosk, enterprise, profile, certificate, VPN, always-on, secure, boot, launcher, accessibility, parental control, screen time, battery saver, low power, overheat, throttle, lag, smooth, fluid, jailbreak, root, custom ROM, bootloader",
         "Pol\u00edtica BYOD (340 w). MCQ.",
         "NIST guide resumida (adaptada).",
         "Email BYOD (60 w)",
         "BYOD: mezcla de datos personal/corporativo"),
        ("Cloud storage basics",
         "How much / how many",
         "+45: cloud, storage, drive, sync, upload, download, backup, restore, share, link, expire, permission, edit, view, comment, owner, collaborator, quota, plan, free, premium, upgrade, bandwidth, latency, region, zone, availability, provider, AWS, Azure, GCP, public, private, hybrid, SLA, uptime, downtime, redundant, encrypted, at rest, in transit, migration, export, import",
         "Comparaci\u00f3n Google Drive / OneDrive / Dropbox (340 w).",
         "VOA tech sobre cloud.",
         "Informe corto (65 w)",
         "Ransomware vs backups inmutables"),
        ("Email etiquette advanced",
         "Polite requests: would you mind \u2026",
         "+35: courteous, polite, professional, formal, informal, clear, concise, brief, detailed, greeting, opening, closing, signature, title, position, request, reminder, follow-up, urgent, high priority, low priority, confidential, sensitive, internal, external, stakeholder, recipient, thread, chain, attachment, link, body, subject line, quote, forward",
         "5 emails con tono variable (350 w).",
         "EFF podcast sobre email encryption (adaptado).",
         "Reescribir email rude (65 w)",
         "BEC: Business Email Compromise"),
        ("Meetings in IT",
         "Used to",
         "+35: agenda, minutes, chair, chairperson, facilitator, attendee, remote, hybrid, in person, video call, conference, bridge, mute, unmute, camera, share screen, record, transcription, breakout, poll, chat, raise hand, round table, standup, retro, retrospective, planning, sprint, demo, review, follow up, action item, blocker, risk",
         "Acta de sprint review (380 w).",
         "Podcast scrum.org (CC BY).",
         "Acta de reuni\u00f3n (70 w)",
         "Grabaciones Zoom: leaks y consentimiento"),
        ("Reporting IT problems",
         "First conditional",
         "+40: incident, ticket, severity, priority, P1, P2, P3, high, medium, low, affects, impact, scope, user, department, root cause, workaround, fix, patch, update, escalate, assign, reassign, close, reopen, SLA, met, missed, response, resolution, categorize, classify, reproduce, confirm, duplicate, unable to reproduce",
         "Reporte de incidente (360 w).",
         "Audio CISA 'incident reporting 101'.",
         "Ticket P2 (70 w)",
         "C\u00f3mo reportar sin filtrar info"),
        ("Cybersecurity \u2013 first terms",
         "Passive voice simple",
         "+45: security, threat, risk, vulnerability, exploit, attack, attacker, hacker, defender, protect, prevent, detect, respond, recover, confidentiality, integrity, availability, CIA triad, asset, data, system, network, endpoint, host, server, perimeter, trust, identity, user, admin, privilege, access, control, policy, rule, audit, log, monitor, alert, alarm, incident",
         "Intro a CIA triad (380 w).",
         "CISA 'What is cybersecurity?' nivel 2.",
         "Definir 5 t\u00e9rminos (70 w)",
         "CIA vs AAA (Authentication, Authorization, Accounting)"),
        ("Phishing awareness",
         "Second conditional b\u00e1sico",
         "+40: phishing, spear phishing, whaling, smishing, vishing, pretexting, baiting, lure, bait, link, attachment, macro, sender, from, reply-to, domain, spoof, impersonation, urgency, fear, reward, click, hover, inspect, report, quarantine, allowlist, denylist, training, simulation, campaign, click rate, report rate, awareness",
         "3 emails phishing reales (FBI PDF, 420 w).",
         "Podcast FBI 'Cyber tips' (adaptado).",
         "Advisory interno (80 w)",
         "C\u00f3mo dise\u00f1ar un phishing test"),
        ("Backup and recovery",
         "Should have (reproche)",
         "+40: backup, full, incremental, differential, snapshot, image, copy, archive, retention, rotate, GFS, 3-2-1, offline, offsite, air gap, immutable, tape, disk, cloud, restore, recovery, RTO, RPO, DR, BCP, test, drill, verify, corrupt, lost, recovered, partial, point in time, freshest, stale, scheduler, job, fail, succeed",
         "Pol\u00edtica 3-2-1 (400 w).",
         "NIST talk adaptado.",
         "Plan de backup (80 w)",
         "Ransomware + backups encriptados"),
        ("Data vs information vs knowledge",
         "Relative pronouns that / which / who",
         "+30: data, raw, processed, clean, dirty, valid, invalid, record, field, table, row, column, dataset, source, pipeline, transform, load, information, report, insight, knowledge, wisdom, action, decide, interpret, analyze, visualize, dashboard, KPI, metric",
         "DIKW pyramid (380 w).",
         "OpenLearn 'Data literacy'.",
         "Explicar DIKW (80 w)",
         "Minimizaci\u00f3n de datos (GDPR)"),
        ("Describing processes",
         "Sequence connectors, passive",
         "+30: step, stage, phase, flow, input, output, trigger, event, action, result, approve, deny, submit, cancel, loop, repeat, until, while, branch, if, else, then, begin, end, start, stop, pause, resume, review, finalize",
         "Diagrama ITIL incident flow (400 w).",
         "ITIL Foundation intro (OpenLearn adaptado).",
         "Describir un proceso (85 w)",
         "Incident response steps"),
        ("Comparing devices",
         "Adjective + enough / too",
         "+25: light, heavy, thin, thick, durable, fragile, portable, rugged, loud, quiet, bright, dim, sharp, blurry, fast, slow, expensive, cheap, affordable, worth, value, pros, cons, spec, benchmark, review",
         "Review de 3 laptops (400 w).",
         "VOA 'Gadget review'.",
         "Review corto (90 w)",
         "TPM chip y secure boot en laptops"),
        ("Writing a tech CV",
         "Present perfect introduction",
         "+30: have worked, have led, have built, have designed, achievement, result, impact, improved, reduced, increased, delivered, launched, saved, optimized, automated, deployed, mentored, trained, certified, accredited, awarded, volunteer, bootcamp, hackathon, open source contributor, maintainer",
         "CV modelo B2/A2 (400 w).",
         "Podcast CC 'CV tips'.",
         "CV de 100 w",
         "Privacidad del CV en ATS"),
        ("Mock KET (A2 exam)",
         "Integrado",
         "1 000 acumuladas",
         "KET Reading completo (~1 200 w, 7 parts).",
         "KET Listening completo (30 min).",
         "2 tasks: mensaje (25 w) + email (50 w)",
         "Reflexi\u00f3n: qu\u00e9 te sorprendi\u00f3")
    ]),

    # ================================================= B1 =================================================
    ("B1", "Intermediate \u2013 PET", "2 000 palabras acumuladas", [
        ("The Software Development Lifecycle",
         "Present perfect continuous",
         "+60: SDLC, requirement, gather, elicit, stakeholder, specification, functional, non-functional, design, high-level, low-level, architecture, component, module, interface, implementation, coding, build, compile, package, artifact, repository, branch, commit, merge, pull request, review, CI, CD, pipeline, test, unit, integration, system, acceptance, regression, smoke, sanity, deploy, staging, production, release, hotfix, rollback, maintenance, retire, sunset, waterfall, V-model, spiral, agile, iterative, incremental, velocity, burndown",
         "'Waterfall vs Agile' (450 w). Matching headings (PET/IELTS 4.5).",
         "IEEE podcast (CC) 'Intro to SDLC' adaptado.",
         "Informe (120 w)",
         "Security in SDLC: SSDLC introducido"),
        ("Programming languages overview",
         "Gerund vs infinitive",
         "+55: Python, Java, C, C++, C#, JavaScript, TypeScript, Go, Rust, Ruby, PHP, Swift, Kotlin, syntax, semantics, compile, interpret, runtime, JIT, garbage collector, memory, pointer, reference, typed, dynamic, static, strong, weak, paradigm, OOP, functional, procedural, script, framework, library, package, module, import, ecosystem, popularity, community, stable, mature, deprecated, versioning, semantic versioning",
         "Tiobe index explicado (480 w). MCQ.",
         "VOA Tech 'Which language should you learn?' nivel 3.",
         "Opini\u00f3n (130 w)",
         "Memory-safe languages (Rust) vs C/C++"),
        ("Databases introduction",
         "Defining relative clauses",
         "+50: database, DBMS, SQL, NoSQL, table, row, record, column, field, primary key, foreign key, index, schema, normal form, denormalize, transaction, ACID, commit, rollback, query, select, insert, update, delete, join, inner, outer, left, right, aggregate, group by, having, view, stored procedure, trigger, backup, restore, replica, master, slave, shard, partition, MongoDB, PostgreSQL, MySQL, Oracle, SQLite",
         "'Relational vs document' (500 w). Summary.",
         "MIT OCW 6.830 intro adaptada.",
         "Comparativa (130 w)",
         "SQL injection intro"),
        ("Cloud computing (IaaS/PaaS/SaaS)",
         "Modals of deduction present",
         "+50: cloud, on-premise, hybrid, multi-cloud, provider, region, zone, VM, container, function, serverless, IaaS, PaaS, SaaS, FaaS, CaaS, pay-as-you-go, billing, cost, reserved, spot, scale up, scale out, elasticity, auto-scale, load balancer, CDN, edge, latency, throughput, tenant, multitenant, isolation, noisy neighbor, shared responsibility model",
         "NIST SP 800-145 adaptado (500 w).",
         "AWS re:Invent keynote corto (adaptado).",
         "Informe decisional (140 w)",
         "Shared responsibility model en seguridad"),
        ("Networks & protocols",
         "Modals of deduction past",
         "+50: TCP, UDP, IP, IPv4, IPv6, packet, datagram, header, payload, port, socket, handshake, SYN, ACK, FIN, reset, MTU, fragment, router, switch, gateway, firewall, proxy, NAT, VLAN, subnet, CIDR, mask, DHCP, DNS, A record, MX, TXT, CNAME, HTTP, HTTPS, TLS, SMTP, IMAP, SSH, FTP, SFTP, MQTT, WebSocket",
         "OSI vs TCP/IP (500 w).",
         "CISA 'Network basics' podcast.",
         "Explicar 3-way handshake (140 w)",
         "SYN flood y DoS"),
        ("Cybersecurity threats overview",
         "Passive with modals",
         "+55: threat, actor, insider, outsider, opportunistic, targeted, nation-state, cybercriminal, script kiddie, hacktivist, motive, financial, political, espionage, sabotage, attack surface, attack vector, kill chain, recon, weaponize, deliver, exploit, install, C2, command and control, exfiltrate, lateral, persistence, privilege escalation, payload, backdoor, rootkit, trojan, worm, virus, ransomware, wiper, spyware, keylogger, adware, cryptojacker, botnet",
         "Cyber kill chain de Lockheed (550 w). T/F/NG.",
         "Podcast CISA 'Threat landscape 2026' adaptado.",
         "Res\u00famen del kill chain (150 w)",
         "MITRE ATT&CK intro"),
        ("Malware types",
         "Reported statements",
         "+45: malware, benign, malicious, payload, dropper, loader, stealer, infostealer, banker, wiper, bot, worm, trojan, virus, ransomware, cryptoware, locker, wiper, RAT, rootkit, bootkit, spyware, adware, keylogger, fileless, living off the land, LOLBAS, PowerShell, WMI, persistence, registry, startup, service, task, DLL sideloading, code signing, obfuscation, packer, unpacker",
         "Caso Emotet (600 w).",
         "Malwarebytes / EFF podcast adaptado.",
         "Clasificar muestras (150 w)",
         "Fileless malware"),
        ("Phishing and social engineering deep dive",
         "Reported questions",
         "+40: pretext, elicit, rapport, trust, urgency, authority, scarcity, reciprocity, consistency, proof, target, profile, dossier, OSINT, doxx, spear, clone, BEC, invoice fraud, gift card scam, MFA fatigue, push bombing, SIM swap, lookalike domain, typosquatting, homoglyph, attachment, macro, RLO, double extension",
         "Anatom\u00eda de un BEC (580 w).",
         "FBI 'Business Email Compromise' podcast.",
         "Advisory (150 w)",
         "Dise\u00f1o de campa\u00f1as de awareness"),
        ("Password management",
         "Verb + -ing / to-infinitive",
         "+35: password, passphrase, entropy, length, complexity, dictionary, rainbow table, brute force, credential stuffing, spray, reuse, rotate, expire, lifetime, history, minimum age, salt, hash, MD5, SHA-1, SHA-256, bcrypt, Argon2, PBKDF2, vault, manager, Bitwarden, KeePass, 1Password, master password, biometric, recovery, seed, TOTP, HOTP",
         "NIST SP 800-63B moderno (620 w).",
         "Lockpicking talk DEFCON (archive.org).",
         "Pol\u00edtica (160 w)",
         "Moderno: no forzar rotaci\u00f3n peri\u00f3dica"),
        ("Multi-factor authentication",
         "Zero conditional",
         "+35: MFA, 2FA, factor, knowledge, possession, inherence, location, behavior, something you know/have/are, SMS, OTP, TOTP, HOTP, push, app, token, hardware, FIDO2, WebAuthn, passkey, authenticator, Authy, Duo, phishable, non-phishable, bypass, fallback, recovery, backup codes, SIM swap, MFA fatigue",
         "Passkeys vs TOTP (600 w).",
         "Yubico / FIDO Alliance talk CC.",
         "Comparativa (160 w)",
         "Caso Uber 2022"),
        ("Encryption basics",
         "Future: will vs going to",
         "+40: encryption, decryption, plaintext, ciphertext, key, algorithm, symmetric, asymmetric, public, private, AES, DES, 3DES, RSA, ECC, Diffie-Hellman, ECDH, hash, SHA-2, SHA-3, digital signature, certificate, CA, PKI, X.509, chain of trust, revocation, CRL, OCSP, HSTS, at rest, in transit, end-to-end, perfect forward secrecy, ephemeral",
         "E2EE en Signal (620 w).",
         "EFF podcast 'E2EE explained'.",
         "Informe para CEO (160 w)",
         "Backdoors vs E2EE: debate"),
        ("Firewalls and antivirus",
         "Defining and non-defining clauses",
         "+35: firewall, stateless, stateful, next-gen, NGFW, WAF, IDS, IPS, rule, allow, deny, drop, reject, log, audit, whitelist, blocklist, signature, heuristic, behavior, sandbox, EDR, XDR, SIEM, SOAR, antivirus, engine, definitions, update, false positive, false negative, quarantine, remediation, rollback, hygiene",
         "EDR vs AV (640 w).",
         "Sophos podcast (CC).",
         "Comparaci\u00f3n (170 w)",
         "Log4Shell: WAFs ante 0-day"),
        ("Ethical hacking introduction",
         "Present perfect vs past simple",
         "+40: ethical, white hat, gray hat, black hat, red team, blue team, purple team, bug bounty, responsible disclosure, CVE, CVSS, exploit, PoC, scope, rules of engagement, engagement letter, NDA, recon, passive, active, OSINT, Shodan, enumeration, Nmap, Nessus, Burp, Metasploit, report, finding, remediation, retest, CTF",
         "'Introduction to pentesting' (640 w).",
         "HackerOne podcast (CC adaptado).",
         "Ethics statement (170 w)",
         "Legalidad en Per\u00fa: Ley 30096"),
        ("IoT devices and risks",
         "Would for hypothetical",
         "+35: IoT, sensor, actuator, gateway, hub, constrained, battery, sleep, duty cycle, MQTT, CoAP, Zigbee, Z-Wave, LoRa, NB-IoT, 5G, firmware, OTA, default password, telemetry, edge, fog, cloud, smart home, smart city, smart meter, industrial, OT, ICS, SCADA, PLC, HMI, air gap, Mirai, botnet, DDoS",
         "Mirai botnet caso (660 w).",
         "Podcast 'Smart Grid security' (OpenLearn).",
         "Pol\u00edtica IoT (180 w)",
         "OT/IT convergence"),
        ("AI & machine learning intro",
         "Unless, as long as",
         "+45: AI, narrow, general, weak, strong, ML, supervised, unsupervised, semi-supervised, reinforcement, label, feature, training, validation, test, split, overfit, underfit, bias, variance, model, algorithm, linear regression, logistic, decision tree, random forest, SVM, KNN, neural network, deep learning, CNN, RNN, transformer, LLM, fine-tune, prompt, inference, latency, throughput",
         "IEEE Spectrum intro a LLMs (680 w).",
         "MIT OCW 6.S897 adaptado.",
         "Resumen (180 w)",
         "Adversarial ML intro"),
        ("Blockchain intro",
         "Third conditional intro",
         "+40: blockchain, ledger, block, transaction, hash, Merkle, nonce, miner, proof of work, proof of stake, consensus, fork, hard fork, soft fork, node, validator, wallet, private key, public key, address, Bitcoin, Ethereum, smart contract, gas, DeFi, DAO, NFT, token, coin, stablecoin, 51% attack, reorg",
         "Caso DAO hack (680 w).",
         "Khan Academy 'Bitcoin' (CC).",
         "Resumen ejecutivo (190 w)",
         "Rug pulls y KYC"),
        ("Project management (Agile/Scrum)",
         "Should / ought to + perfect",
         "+40: agile, manifesto, scrum, sprint, epic, story, task, subtask, backlog, refinement, grooming, planning, daily, review, retro, done, definition of done, acceptance criteria, velocity, burndown, burnup, kanban, WIP limit, flow, cycle time, lead time, blocker, impediment, stakeholder, PO, SM, dev team, T-shirt sizing, story points",
         "'Scrum in 10 minutes' (700 w).",
         "Spotify engineering podcast (CC).",
         "Reporte de sprint (200 w)",
         "Security stories en el backlog"),
        ("Writing technical emails",
         "Linking words: however, moreover, therefore",
         "+30: escalation, approval, sign-off, confidential, restricted, internal use, external, FYI, EOM, TBD, TL;DR, follow up, loop in, ping, raise, flag, clarify, confirm, attach, enclosed, forwarded, revised, updated, deprecated, pending, on hold, resumed, completed, accepted",
         "3 emails con registros distintos (720 w).",
         "'Email writing' OpenLearn.",
         "Email B1 140 w",
         "Marcado 'Confidential' \u2013 DLP"),
        ("Giving a short technical talk",
         "Discourse markers",
         "+25: introduce, outline, agenda, overview, motivate, problem, challenge, hypothesis, approach, method, result, conclusion, takeaway, Q&A, summarize, recap, transition, moving on, to sum up, in conclusion, any questions, floor is yours, thank you for your attention, feel free to interrupt",
         "Estructura de charla (700 w).",
         "TED-Ed 'Speak like a pro'.",
         "Outline (180 w)",
         "Talks sensibles: info classification"),
        ("Mock PET + B1 project",
         "Integrado",
         "2 000 acumuladas",
         "PET Reading completo (~2 800 w).",
         "PET Listening completo (35 min).",
         "2 tasks: email (100 w) + article (100 w)",
         "Cierre: self-assessment con CEFR grid"),
    ]),

    # ================================================= B2 =================================================
    ("B2", "Upper-Intermediate \u2013 FCE", "4 000 palabras acumuladas", [
        ("Cybersecurity incident response",
         "Mixed conditionals",
         "+80: preparation, identification, containment, eradication, recovery, lessons learned, playbook, runbook, CSIRT, on-call, rotation, triage, severity, scope, IOC, TTP, forensic, memory, disk, network, triage, timeline, correlation, normalize, enrich, artifact, chain of custody, legal hold, communication, stakeholder, exec, PR, legal, regulator, disclosure, notification, deadline, 72-hour, containment strategy, isolate, quarantine, block, kill switch",
         "NIST SP 800-61r2 adaptado (800 w). Multiple matching.",
         "SANS webinar 'IR lifecycle' (adaptado CC).",
         "IR plan (180 w)",
         "SolarWinds case study"),
        ("Network security architecture",
         "Cleft sentences",
         "+70: defense in depth, perimeter, zone, DMZ, segment, microsegment, east-west, north-south, least privilege, zero trust, SD-WAN, SASE, ZTNA, CASB, SWG, proxy forward, reverse, IDS, IPS, NDR, TAP, SPAN, flow, NetFlow, sFlow, IPFIX, packet capture, Wireshark, pcap, deep packet inspection, TLS inspection, MITM, PKI internal, certificate pinning",
         "Secure reference architecture (820 w).",
         "CISA webinar 'Zero trust basics'.",
         "Architecture memo (200 w)",
         "TLS inspection vs privacy"),
        ("Advanced malware & ransomware",
         "Emphatic structures: do / does",
         "+65: ransomware as a service, RaaS, affiliate, broker, initial access broker, IAB, double extortion, triple, leak site, data auction, negotiator, wallet, ransom note, encrypt, extort, exfiltrate, wiper, wiperware, LockBit, Conti, ALPHV, BlackCat, Play, residential, business, critical infra, OT, factory, hospital, pipeline, payment, Bitcoin, Monero, OFAC, sanction",
         "Caso Colonial Pipeline (850 w).",
         "Podcast Darknet Diaries (algunos eps CC).",
         "Board memo (200 w)",
         "Pagar o no pagar: el dilema"),
        ("Zero-day vulnerabilities",
         "Inversion after negative adverbs",
         "+55: zero-day, N-day, unknown, unpatched, patch gap, exploitation in the wild, ITW, PoC, weaponized, broker, ZDI, Pwn2Own, bug bounty, CVE, CVSS, EPSS, exploit kit, drive-by, watering hole, supply chain, hypervisor, kernel, userland, sandbox escape, chain, RCE, LPE, info leak, memory corruption, use-after-free, heap, stack, overflow",
         "'Anatomy of a zero-day' (880 w).",
         "Google Project Zero talk (CC).",
         "Risk advisory (220 w)",
         "Disclosure ethics"),
        ("OWASP Top 10 web app security",
         "Concessive clauses: despite, although",
         "+70: OWASP, A01 broken access control, A02 cryptographic failure, A03 injection, A04 insecure design, A05 misconfig, A06 outdated, A07 auth failure, A08 integrity, A09 logging, A10 SSRF, IDOR, BOLA, BOPLA, mass assignment, CSRF, XSS, stored, reflected, DOM, prototype pollution, SSTI, XXE, deserialization, JWT, hardcoded",
         "Top 10 2025 draft (900 w).",
         "OWASP Global AppSec talk (CC).",
         "Threat model (220 w)",
         "AppSec en CI/CD"),
        ("Cryptography fundamentals",
         "Participle clauses",
         "+60: cryptosystem, primitive, cipher, block, stream, mode, ECB, CBC, GCM, CTR, nonce, IV, tag, authenticated encryption, AEAD, KDF, HKDF, scrypt, Argon2, MAC, HMAC, PRF, PRNG, CSPRNG, entropy, seed, side channel, timing, padding oracle, Bleichenbacher, downgrade, BEAST, POODLE, Heartbleed, Lucky13",
         "Why 'roll your own crypto' fails (900 w).",
         "Charla Cryptographers' Panel RSA (archive).",
         "Policy memo (220 w)",
         "FIPS 140-3"),
        ("Public Key Infrastructure (PKI)",
         "Nominalization",
         "+50: PKI, CA, intermediate, root, trust store, certificate, CSR, CN, SAN, SNI, EV, DV, OV, chain, revocation, CRL, OCSP, OCSP stapling, CT log, Certificate Transparency, pinning, HSTS, preload, wildcard, SAN multi-domain, SCT, issuance, renewal, ACME, Let's Encrypt, HSM, key ceremony, offline root",
         "'CT and CA compromises' (920 w).",
         "Let's Encrypt podcast (CC).",
         "Plan PKI interno (230 w)",
         "Internal CA vs public"),
        ("VPNs & secure communication",
         "Wish + past perfect",
         "+45: VPN, site-to-site, remote access, split tunneling, full tunnel, IPSec, IKEv2, OpenVPN, WireGuard, TLS 1.3, mTLS, DTLS, QUIC, HTTP/3, DoH, DoT, eSNI, ECH, meshVPN, Tailscale, Nebula, ZTNA, hairpin, kill switch, DNS leak, IPv6 leak, provider, no-log, audit",
         "WireGuard vs IPSec (920 w).",
         "WireGuard paper talk (CC).",
         "Decisi\u00f3n t\u00e9cnica (230 w)",
         "VPN en pa\u00edses con censura"),
        ("SIEM and log analysis",
         "Passive reporting: it is said that",
         "+50: SIEM, log source, parser, normalization, CIM, ECS, index, search, query, KQL, SPL, dashboard, detection, rule, correlation, alert, alarm, ticket, MTTR, MTTD, tiering, T1, T2, T3, analyst, hunter, playbook, SOAR, orchestration, automation, response, false positive, false negative, tuning, threshold, Sigma",
         "'Building a detection pipeline' (950 w).",
         "Splunk .conf keynote (adaptado CC).",
         "Detection design (240 w)",
         "Logs como evidencia forense"),
        ("Penetration testing",
         "Causative have / get",
         "+45: scope, rules of engagement, black box, gray box, white box, PTES, PTS, methodology, reconnaissance, scanning, enumeration, exploitation, privilege escalation, lateral movement, persistence, data exfiltration, cleanup, reporting, CVSS, CWE, OWASP WSTG, NIST SP 800-115, internal, external, wireless, physical, social, web app, API, cloud",
         "Sample pentest report redacted (980 w).",
         "Offensive Security podcast (CC).",
         "Executive summary (240 w)",
         "Retesting and remediation"),
        ("AI & cybersecurity",
         "Subjunctive after suggest / recommend",
         "+60: AI/ML in security, anomaly detection, clustering, baseline, drift, concept drift, adversarial ML, evasion, poisoning, membership inference, model extraction, prompt injection, jailbreak, guardrail, red teaming AI, LLM security, OWASP LLM Top 10, RAG leaks, data leakage, hallucination risk, synthetic content, deepfake, voice clone, AI-generated phishing",
         "OWASP LLM Top 10 (980 w).",
         "Anthropic / NIST AI RMF podcast (CC/PD).",
         "AI policy draft (240 w)",
         "Prompt injection defense"),
        ("Cloud security",
         "Nominal clauses",
         "+55: CSPM, CNAPP, CWPP, CIEM, CASB, IAM, SCIM, SAML, OIDC, SSO, least privilege, just-in-time, PIM, PAM, service account, managed identity, secret, rotation, KMS, HSM, envelope encryption, bucket, public access, S3, Blob, GCS, data exfil, IMDS, IMDSv2, SSRF to metadata, landing zone, guardrail, policy as code",
         "Caso Capital One 2019 (1 000 w).",
         "AWS re:Inforce talk (adaptado CC).",
         "Runbook CSPM (250 w)",
         "Shared fate model"),
        ("DevSecOps",
         "Subordinate clauses complejas",
         "+50: shift left, shift right, SAST, DAST, IAST, SCA, secret scanning, license scanning, IaC scanning, tfsec, checkov, policy as code, OPA, Rego, Sentinel, signing, Sigstore, cosign, SBOM, CycloneDX, SPDX, provenance, SLSA, supply chain, attestation, pipeline security, runner, artifact, registry",
         "'SLSA levels explained' (1 020 w).",
         "CNCF sigstore talk (CC).",
         "SDL policy (250 w)",
         "XZ backdoor 2024"),
        ("IoT security",
         "Discourse markers formales",
         "+40: attack surface, firmware, secure boot, measured boot, chain of trust, TPM, TEE, trust zone, TrustZone, ARM, Intel SGX, over-the-air, OTA, secure channel, device identity, DICE, PKI for IoT, gateway, edge, zero touch provisioning, BOM, hardware supply chain, tampering, glitching, fault injection, side channel, DPA",
         "'Securing constrained devices' (1 050 w).",
         "IEEE IoT talk (CC).",
         "IoT security brief (260 w)",
         "Mirai revisited"),
        ("Mobile security",
         "Inversion for emphasis",
         "+45: Android security model, SELinux, Verified Boot, Play Protect, sideloading, permission, runtime, scoped storage, iOS security, Secure Enclave, data protection classes, MDM, MAM, supervised mode, DEP, ABM, jailbreak, root, Frida, Magisk, reverse engineering, APK, IPA, Android WebView, deep link, URL scheme, universal link",
         "'iOS vs Android threat models' (1 050 w).",
         "OWASP MAS talk (CC).",
         "BYOD policy v2 (260 w)",
         "Pegasus y mercenary spyware"),
        ("GDPR & data privacy",
         "Modal verbs in legal register",
         "+40: GDPR, CCPA, LGPD, PIPL, PIPEDA, Peru Ley 29733, data subject, controller, processor, DPO, DPIA, ROPA, lawful basis, consent, legitimate interest, necessity, purpose limitation, data minimization, storage limitation, accuracy, integrity, accountability, DSR, right of access, erasure, rectification, portability, objection, fine, 4 percent",
         "GDPR simplified (1 080 w).",
         "EDPB webinar (CC).",
         "DPIA template (280 w)",
         "Transferencias internacionales post-Schrems II"),
        ("Writing a security report",
         "Hedging language",
         "+30: appears, suggests, consistent with, likely, unlikely, probable, possible, confidence level, low, medium, high, assessment, scoping, methodology, findings, severity, likelihood, impact, CVSS, remediation, mitigation, residual risk, recommendation, timeline, short-term, medium-term, long-term, appendix, annex, glossary",
         "Template Mandiant-style (1 100 w).",
         "Mandiant reports walkthrough (CC).",
         "Informe a CISO (300 w)",
         "Confidence levels (ICD 203)"),
        ("Presenting security findings",
         "Rhetorical questions",
         "+25: takeaway, headline finding, so what, impact statement, business impact, regulatory impact, reputational, operational, financial, quantify, estimate, conservative, best case, worst case, compensating control, accepted risk, risk owner, risk register, treatment plan",
         "Board deck modelo (1 100 w).",
         "Gartner briefing podcast (CC).",
         "Elevator pitch 200 w",
         "Board-level language"),
        ("Debate: privacy vs security",
         "Concession and counter-argument",
         "+30: on the one hand, on the other hand, admittedly, granted, nevertheless, nonetheless, having said that, that said, not to mention, moreover, furthermore, in addition, by contrast, conversely, in sharp contrast, weigh, trade-off, balance, proportionate, necessary, sufficient",
         "Apple vs FBI 2016 revisitado (1 120 w).",
         "Berkman Klein debate (CC BY).",
         "Argumentative (280 w)",
         "Client-side scanning (CSAM)"),
        ("Mock FCE + B2 portfolio",
         "Integrado",
         "4 000 acumuladas",
         "FCE Reading & UoE completo (~3 500 w).",
         "FCE Listening completo (40 min).",
         "2 tasks: essay (140\u2013190 w) + review/report (140\u2013190 w)",
         "Entregable: portafolio de incidente"),
    ]),

    # ================================================= C1 =================================================
    ("C1", "Advanced \u2013 CAE", "8 000 palabras acumuladas", [
        ("Advanced threat intelligence",
         "Inversion for formal register",
         "+120: CTI, strategic, operational, tactical, technical, stakeholder, consumer, producer, IOC, IOA, TTP, diamond model, adversary, infrastructure, capability, victim, pivot, enrichment, attribution, confidence, ICD 203, analytic tradecraft, competing hypotheses, ACH, cognitive bias, anchoring, confirmation, collection, requirements, PIR, SIR, source, primary, secondary, open, closed, dark web, paste site, forum, Telegram, cybercrime ecosystem, feed, TAXII, STIX",
         "Podes. Vol. de CrowdStrike Global Threat Report adaptado (1 250 w).",
         "Keynote RSA CTI panel (CC).",
         "Intel assessment (250 w)",
         "Attribution pitfalls"),
        ("APTs (Advanced Persistent Threats)",
         "Cleft + pseudo-cleft",
         "+110: APT1, APT28, APT29, Lazarus, Equation Group, Turla, Sandworm, Volt Typhoon, Flax Typhoon, Mustang Panda, Kimsuky, lifecycle, long dwell time, living off the land, LOLBAS, GTFOBins, fileless, BYOVD, driver, kernel, rootkit, bootkit, UEFI, firmware implants, MoonBounce, BlackLotus, staging, ex\u00adfil, C2 domain, DGA, fast flux, domain shadowing, front, TLS cert rotation",
         "Volt Typhoon CISA advisory (1 250 w).",
         "FBI brief + Mandiant podcast (CC/PD).",
         "Hunt plan (260 w)",
         "Prepositioning en OT"),
        ("Zero Trust Architecture",
         "Subjunctive for recommendations",
         "+95: NIST SP 800-207, policy decision point, policy enforcement point, PE, PEP, PDP, PA, subject, resource, trust algorithm, score-based, criteria-based, explicit verify, continuous verification, BeyondCorp, Google, identity-aware proxy, context-aware access, device posture, attestation, EDR signal, risk score, ITDR, conditional access, session, token binding, proof of possession, DPoP, mTLS",
         "NIST SP 800-207 completo adaptado (1 300 w).",
         "Google BeyondCorp talk (CC).",
         "Implementation roadmap (260 w)",
         "ZT en OT/ICS"),
        ("SOC operations",
         "Advanced passive forms",
         "+90: SOC, MSSP, MDR, XDR, SOAR, tier, L1, L2, L3, escalation, shift handover, runbook, playbook, SLA, MTTA, MTTD, MTTR, ticketing, queue, backlog, KPIs, SLIs, burnout, staffing model, follow-the-sun, CSIRT, CERT, ISAC, ISAO, information sharing, TLP, WHITE, GREEN, AMBER, AMBER+STRICT, RED, CLEAR",
         "'Inside a modern SOC' (1 320 w).",
         "Black Hat SOC panel (CC).",
         "SOC design doc (270 w)",
         "Burnout y retenci\u00f3n"),
        ("Threat hunting methodologies",
         "Complex comparatives",
         "+80: hypothesis-driven, data-driven, intel-driven, situational awareness, diamond of intrusion, pyramid of pain, hash, IP, domain, artifact, network, tools, TTPs, Mandiant maturity model, MaGMa, TTPs over IOCs, Cobalt Strike, Metasploit, Sliver, BruteRatel, Havoc, Empire, PowerSploit, artifacts, Sysmon, ETW, osquery, Velociraptor, GRR",
         "'Hunting Cobalt Strike beacons' (1 350 w).",
         "SANS hunt talk (CC).",
         "Hunt report (280 w)",
         "Beacon jitter & sleep"),
        ("Malware reverse engineering",
         "Nominal groups densos",
         "+90: static analysis, dynamic analysis, disassembler, decompiler, IDA Pro, Ghidra, Binary Ninja, radare2, Cutter, x64dbg, OllyDbg, WinDbg, Frida, debugger, breakpoint, watchpoint, memory dump, process injection, reflective, DLL, APC, early bird, process hollowing, atom bombing, syscall, direct, indirect, unhooking, EDR evasion, packer, UPX, Themida, VMProtect, anti-debug, anti-vm, anti-analysis",
         "'Reverse engineering a dropper' (1 380 w).",
         "REcon talk (archive CC).",
         "RE report (300 w)",
         "Legalidad del RE en Per\u00fa"),
        ("Machine learning in cybersecurity",
         "Participle clauses complejas",
         "+85: feature engineering, tabular, image, text, embedding, sequence, temporal, autoencoder, isolation forest, DBSCAN, GMM, HMM, LSTM, GRU, transformer, attention, explainable AI, SHAP, LIME, counterfactual, drift, concept drift, data drift, monitoring, feedback loop, human in the loop, active learning, labeling, weak supervision",
         "arXiv paper sobre ML-based IDS (1 400 w).",
         "MIT Lincoln Labs talk (CC).",
         "Design doc ML-IDS (300 w)",
         "Baseline bias"),
        ("Adversarial AI attacks",
         "Modal perfect continuous",
         "+75: adversarial example, perturbation, epsilon, FGSM, PGD, Carlini-Wagner, DeepFool, transferability, black box, white box, query-based, gradient-based, evasion, poisoning, backdoor, trojaning, clean label, BadNets, model stealing, membership inference, model inversion, watermark, canary, defensive distillation, adversarial training, certified robustness",
         "NIST AI 100-2 adaptado (1 420 w).",
         "ICLR talk (CC).",
         "Risk matrix (310 w)",
         "Jailbreaks in LLMs"),
        ("Quantum computing threats",
         "Advanced hedging",
         "+70: qubit, superposition, entanglement, decoherence, gate, quantum volume, logical qubit, error correction, surface code, Shor's algorithm, Grover's, RSA, ECC, symmetric, hash, harvest now decrypt later, HNDL, store now decrypt later, Q-day, cryptographically relevant quantum computer, CRQC, roadmap, IBM, Google, IonQ, Rigetti, D-Wave, NISQ",
         "IBM quantum roadmap (1 440 w).",
         "IBM research podcast (CC).",
         "Risk brief (320 w)",
         "HNDL timeline"),
        ("Post-quantum cryptography",
         "Inversion + concession",
         "+70: PQC, lattice-based, code-based, multivariate, hash-based, isogeny, CRYSTALS-Kyber, ML-KEM, CRYSTALS-Dilithium, ML-DSA, SPHINCS+, SLH-DSA, FALCON, FIPS 203, 204, 205, 206, hybrid, composite, migration, crypto agility, inventory, CBOM, cryptographic bill of materials, NIST PQC project, round 4, SIKE broken",
         "FIPS 203 ML-KEM (1 450 w).",
         "NIST PQC webinar (PD).",
         "Migration plan (330 w)",
         "SIKE break lesson"),
        ("Critical infrastructure protection",
         "Formal cohesion",
         "+65: CIP, CIKR, NIS2, Executive Order 13636, CISA sectors, energy, water, transportation, healthcare, financial, food, emergency, OT, ICS, SCADA, DCS, PLC, RTU, HMI, Purdue model, level 0-5, air gap, data diode, one-way, unidirectional, zone, conduit, IEC 62443, NERC CIP",
         "NIS2 Directive summary (1 460 w).",
         "ENISA podcast (CC).",
         "Gap analysis (340 w)",
         "Stuxnet lessons"),
        ("Supply chain attacks",
         "Cohesive devices avanzados",
         "+70: SolarWinds, Kaseya, XZ Utils, 3CX, CodeCov, event-stream, ua-parser-js, dependency confusion, typosquatting, brandjacking, malicious package, npm, PyPI, RubyGems, Go modules, crates.io, Maven, NuGet, SBOM, VEX, CycloneDX, SPDX, SLSA, provenance, attestation, reproducible builds, Sigstore, cosign, witness",
         "XZ Utils backdoor anatomy (1 480 w).",
         "CISA SBOM talk (PD).",
         "SBOM requirements (340 w)",
         "Open source maintainer exhaustion"),
        ("Dark web & cybercrime",
         "Metaphor and register",
         "+60: dark web, Tor, hidden service, onion, I2P, Freenet, marketplace, RaaS, phishing kit, checker, stealer log, combo list, credential stuffing, drop, mule, cashout, BIN, dump, skimmer, Magecart, formjacking, triangulation fraud, bulletproof hosting, fast flux, money mule, mixer, tumbler, OFAC, chain analysis, Chainalysis",
         "Europol dark markets report (1 500 w).",
         "BBC podcast 'Cybercrime' (ciertos CC).",
         "Brief for LE (350 w)",
         "Chain analysis in Monero"),
        ("Digital forensics",
         "Technical hedging",
         "+65: DFIR, disk image, write blocker, hash, MD5, SHA-256, volatility, memory, plaster, Rekall, Volatility 3, shimcache, amcache, prefetch, MFT, USN journal, registry hive, shellbags, recent docs, jump list, timeline, supertimeline, plaso, log2timeline, artifacts, KAPE, triage, EDR artifacts, BitLocker, FileVault, TPM, key extraction, custody",
         "'Timeline analysis' (1 520 w).",
         "SANS DFIR summit (CC).",
         "Forensic report (360 w)",
         "Admissibility of digital evidence"),
        ("Cloud-native security (Kubernetes)",
         "Complex noun phrases",
         "+70: K8s, control plane, etcd, API server, scheduler, controller manager, kubelet, kube-proxy, CRI, CNI, CSI, runtime, containerd, CRI-O, pod security, admission, OPA Gatekeeper, Kyverno, Falco, Tetragon, eBPF, network policy, Cilium, service mesh, Istio, Linkerd, mTLS, RBAC, ABAC, service account, tokens, audit log, bench, kube-bench, kube-hunter",
         "Kubernetes threat matrix (1 540 w).",
         "KubeCon talk (CC).",
         "Hardening guide (360 w)",
         "eBPF for security"),
        ("Red team vs Blue team",
         "Subtle register contrast",
         "+55: red team, blue team, purple team, assume breach, adversary emulation, MITRE Caldera, Atomic Red Team, detection engineering, Sigma, content, rule, coverage, heatmap, maturity, CMMC, NIST CSF, ISO 27001, SOC2, tabletop, crisis exercise, war game, rules of engagement, trust agreement, scoping, deconfliction, legal authorization, sign-off",
         "'Purple teaming that works' (1 560 w).",
         "SANS purple team summit (CC).",
         "Exercise plan (370 w)",
         "Legal authorization docs"),
        ("Writing white papers",
         "Academic register",
         "+50: abstract, executive summary, introduction, background, related work, methodology, approach, architecture, findings, discussion, limitation, future work, conclusion, appendix, bibliography, IEEE citation, ACM citation, copyedit, peer review, preprint, arXiv, camera-ready, figure, caption, table, equation, algorithm pseudocode",
         "IEEE white paper template (1 600 w).",
         "ACM keynote on writing.",
         "White paper section (380 w)",
         "Double-blind vs single-blind"),
        ("Academic research reading",
         "Dense academic syntax",
         "+55: hypothesis, research question, dataset, cohort, control, ablation, ablate, benchmark, SOTA, reproduce, reproducibility crisis, code release, artifact evaluation, significance, p-value, effect size, confidence interval, false discovery, multiple testing, Bonferroni, peer review, double blind, reviewer 2, meta-analysis",
         "arXiv cybersecurity paper (1 650 w).",
         "NeurIPS talk (CC).",
         "Literature review (390 w)",
         "Reproducibility in security"),
        ("Panel discussions & negotiation",
         "Turn-taking in English",
         "+40: if I may, to build on that, picking up on, let me jump in, with respect, counterpoint, on that note, indeed, precisely, I take your point, however, push back, challenge, concede, yield the floor, bottom line, wrap up, parking lot",
         "Transcript BlackHat panel (1 680 w).",
         "DEF CON panel (CC).",
         "Opening statement (400 w)",
         "Negotiating disclosure"),
        ("Mock CAE + C1 capstone",
         "Integrado",
         "8 000 acumuladas",
         "CAE Reading & UoE (~4 500 w).",
         "CAE Listening (40 min).",
         "Essay (220\u2013260 w) + proposal/report (220\u2013260 w)",
         "Capstone: threat model completo"),
    ]),

    # ================================================= C2 =================================================
    ("C2", "Proficiency \u2013 CPE", "16 000+ palabras acumuladas", [
        ("Cyber warfare & geopolitics",
         "Nominalization densa y registro acad\u00e9mico",
         "+250: statecraft, deterrence, escalation, compellence, proxy, hybrid warfare, grey zone, below the threshold, attribution, plausible deniability, norms, Tallinn Manual, UN GGE, OEWG, sovereignty, non-intervention, due diligence, state responsibility, countermeasures, necessity, proportionality, armed attack, Article 51, use of force, jus ad bellum, jus in bello, law of armed conflict, LOAC, IHL, distinction, proportionality, dual-use, critical infrastructure, election interference, influence operation, information operation, IO, cognitive warfare, disinformation",
         "Tallinn Manual 2.0 extract (1 900 w). Summary + compare.",
         "Harvard Berkman Klein panel (CC BY).",
         "Policy essay (320 w)",
         "Ukraine conflict cyber dimension"),
        ("State-sponsored actors",
         "Modal perfect + concession",
         "+200: Fancy Bear, Cozy Bear, Lazarus, Kimsuky, APT40, APT41, Gamaredon, Sandworm, Volt Typhoon, Salt Typhoon, Flax Typhoon, doxxing of operators, indictment, sanctions, Treasury OFAC, reward for justice, operator handles, double-dipping, moonlighting, contractor, front company, shell, procurement chain, proliferation, commercial spyware, mercenary, NSO, Candiru, Cytrox, Intellexa, zero-click, forensic triage",
         "Reuters investigative piece adaptado (1 950 w).",
         "ThreatConnect podcast (CC).",
         "Attribution memo (340 w)",
         "Mercenary spyware marketplace"),
        ("International cyber law",
         "Legal English avanzado",
         "+180: Budapest Convention, Additional Protocol II, MLAT, CLOUD Act, E-evidence, lawful access, jurisdiction, extraterritoriality, long-arm statute, GDPR Chapter V, SCC, BCR, adequacy, Schrems I, Schrems II, EU-US Data Privacy Framework, PCLOB, ECHR, Section 702, FISA, Executive Order 14086, legal basis, judicial redress",
         "Schrems II decision summary (1 980 w).",
         "EDPS podcast (CC).",
         "Legal memo (360 w)",
         "Cloud Act vs GDPR"),
        ("Offensive security ethics",
         "Subjunctive in formal prose",
         "+150: dual-use, Wassenaar, EAR, export control, intrusion software, end-use, end-user, mercenary, responsible disclosure, coordinated disclosure, CVD, exceptions, CFAA, Computer Misuse Act, Ley 30096 Peru, good-faith security research, safe harbor, bug bounty scope, gag order, NDA, chilling effect, cooperative disclosure, ZDI, VINCE",
         "DOJ CFAA policy 2022 adaptado (1 990 w).",
         "EFF podcast on security research.",
         "Ethics whitepaper (380 w)",
         "Research vs research law"),
        ("Post-quantum cryptography research",
         "Mathematical English register",
         "+170: lattice, LWE, RLWE, MLWE, SIS, MSIS, Kyber, Dilithium, FALCON, NTRU, NTRU Prime, Saber, hybrid schemes, Kyber+X25519, classic McEliece, BIKE, HQC, HQC-KEM, SPHINCS+, SLH-DSA, XMSS, LMS, stateful hash-based, side-channel, constant time, masking, shuffling, fault attack, certified implementation",
         "arXiv PQC paper (2 050 w).",
         "NIST PQC workshop (PD).",
         "Technical review (400 w)",
         "Side channel in Kyber impls"),
        ("Homomorphic encryption",
         "Densidad t\u00e9rmica",
         "+140: FHE, partial HE, somewhat HE, PHE, SHE, BFV, BGV, CKKS, TFHE, FHEW, bootstrapping, key switching, relinearization, modulus switching, ciphertext slot, SIMD packing, noise budget, multiplicative depth, circuit, leveled, unbounded, OpenFHE, Microsoft SEAL, Lattigo, Concrete, private inference, federated, MPC",
         "OpenFHE whitepaper (2 080 w).",
         "IBM HE talk (CC).",
         "Use-case memo (410 w)",
         "Private ML inference"),
        ("Zero-knowledge proofs",
         "Complex subordination",
         "+130: ZKP, ZK-SNARK, ZK-STARK, Bulletproofs, PLONK, Halo, Groth16, transparent setup, trusted setup, ceremony, toxic waste, prover, verifier, witness, statement, soundness, completeness, zero-knowledge, succinctness, non-interactivity, commitment, polynomial commitment, KZG, IPA, arithmetic circuit, R1CS, Plonkish, zkVM, zkEVM, ZK rollup",
         "ZK-SNARKs paper adaptado (2 100 w).",
         "ZK Summit talk (CC).",
         "Architecture brief (420 w)",
         "ZK in identity (zk-passport)"),
        ("Critical analysis of research papers",
         "Critical-academic register",
         "+110: strength, weakness, threats to validity, internal validity, external validity, construct validity, statistical validity, confounder, selection bias, sampling, representativeness, ablation, baseline, gold standard, ground truth, data leakage, train-test leakage, overclaim, unfalsifiable, cherry-picked, p-hacking, HARKing, preregistration, replication",
         "IEEE S&P paper (2 150 w) + shadow PC review.",
         "USENIX Security panel (CC).",
         "Peer review letter (440 w)",
         "Rebuttal etiquette"),
        ("Peer review process",
         "Metadiscourse",
         "+90: shepherd, AE, area chair, PC chair, senior PC, ACs, reviewer, meta-review, bid, assignment, TPMS, conflict of interest, revision, major revision, minor revision, reject and resubmit, artifact evaluation, camera-ready, reproducibility badge, shadow PC, tutorial, workshop, co-location",
         "USENIX reviewer guide (2 180 w).",
         "Academic leadership panel (CC).",
         "Review (440 w)",
         "Anonymity & responsible review"),
        ("Technical specifications (RFC reading)",
         "Normative prose",
         "+100: MUST, MUST NOT, SHOULD, SHOULD NOT, MAY, RECOMMENDED, OPTIONAL, normative, informative, wire format, state machine, error code, extension, negotiation, backwards compatibility, forward compatibility, deprecation, IANA registry, code point, algorithm identifier, ASN.1, DER, BER, CBOR, COSE, JOSE, JWS, JWE, JWT",
         "RFC 9000 (QUIC) secci\u00f3n clave (2 200 w).",
         "IETF 119 session (CC).",
         "Interop analysis (460 w)",
         "RFC \u00e9tica y seguridad"),
        ("Standards bodies",
         "Precise register",
         "+80: NIST, ISO, IEEE, IEC, ETSI, IETF, W3C, OASIS, TCG, FIDO Alliance, PCI SSC, ENISA, NCSC UK, BSI Germany, ANSSI France, SAI, CERT/CC, ISA, 62443, CIS, CVE MITRE, NVD, JPCERT, standardization process, public comment, call for papers, working group, plenary, consensus, rough consensus, running code",
         "ISO 27001:2022 mapping (2 220 w).",
         "ISO/IEC JTC 1 SC27 podcast (CC).",
         "Standards mapping (460 w)",
         "NIST CSF 2.0"),
        ("Conference presentation skills",
         "Oratory register",
         "+60: hook, signpost, chunk, pause, prosody, stress, emphasis, speak to the slide, speak past the slide, audience engagement, live demo, fail-safe demo, backup recording, stage presence, vocal variety, pacing, Q&A, hostile question, deflection, bridge, hard pivot, stage management, AV, lav mic, wireless, clicker, mirror, Keynote, Reveal.js",
         "Transcripts USENIX Security lightning talks (2 230 w).",
         "BlackHat keynote (CC).",
         "Talk outline (470 w)",
         "Disclosing 0-day live: etiquette"),
        ("Writing for top-tier journals",
         "Academic writing culture",
         "+80: IEEE Transactions, ACM TOPS, TISSEC, CCS, S&P, USENIX, NDSS, DIMVA, RAID, ESORICS, cover letter, letter to editor, desk reject, rolling review, impact factor, H-index, citation, LaTeX, Overleaf, BibTeX, ACM SIG, camera-ready, supplementary material, artifact, dataset DOI, ORCID",
         "ACM CCS 'how to write' (2 260 w).",
         "Editors panel (CC).",
         "Paper draft section (480 w)",
         "Data availability statements"),
        ("Grant and proposal writing",
         "Bureaucratic English",
         "+70: call for proposals, CFP, RFP, LOI, letter of intent, preproposal, full proposal, aims, significance, innovation, approach, broader impacts, intellectual merit, timeline, milestones, deliverables, risks, mitigation, budget, direct cost, indirect, overhead, PI, Co-PI, collaborator, letter of support, subaward, NSF, NIH, DARPA, EU Horizon, ERC, FWF, FAPESP, CONCYTEC",
         "DARPA BAA adaptado (2 290 w).",
         "NSF PI panel (PD).",
         "Proposal aims section (480 w)",
         "CONCYTEC calls for cyber"),
        ("AI safety & alignment",
         "Academic philosophy register",
         "+70: alignment, specification, robustness, assurance, evaluation, red teaming, dangerous capability, autonomous replication, CBRN uplift, persuasion, deception, sycophancy, scheming, reward hacking, goal misgeneralization, inner alignment, outer alignment, interpretability, mechanistic, activation patching, probing, RLHF, DPO, constitutional AI, AISI, NIST AI 100-2, EU AI Act, frontier model",
         "Anthropic RSP adaptado (2 300 w).",
         "AI Safety Summit panel (CC).",
         "Position paper (490 w)",
         "Model welfare"),
        ("Deepfakes & disinformation",
         "Social-science register",
         "+60: deepfake, shallowfake, cheapfake, voice clone, lip sync, reenactment, provenance, C2PA, content credentials, watermark, fingerprint, detection, adversarial robustness, liar's dividend, illusory truth, inoculation, prebunking, debunking, platform policy, takedown, coordinated inauthentic behavior, CIB, influence operation, astroturf",
         "Microsoft-Adobe C2PA whitepaper (2 330 w).",
         "Stanford Internet Observatory (CC).",
         "Platform policy (500 w)",
         "Election integrity 2026"),
        ("Biometric security futures",
         "Legal + technical fusion",
         "+60: biometric, fingerprint, minutiae, iris, retina, face, voice, gait, keystroke, behavior, liveness, presentation attack, PAD, ISO/IEC 30107, template protection, cancelable, fuzzy vault, fuzzy commitment, homomorphic, federated, on-device, secure enclave, FIDO biometric, GDPR Article 9, BIPA, Clearview AI",
         "NIST FRVT report (2 340 w).",
         "FIDO Alliance webinar (CC).",
         "Ethics brief (510 w)",
         "Algorithmic bias y demograf\u00eda"),
        ("Neuromorphic and novel computing",
         "Research frontier register",
         "+60: neuromorphic, spiking neural network, SNN, memristor, resistive RAM, phase change memory, in-memory computing, analog, photonic, silicon photonics, quantum photonic, silicon qubit, topological qubit, Majorana, adiabatic, annealing, D-Wave, superconducting, transmon, trapped ion, neutral atom, quantum internet, QKD, decoherence time",
         "IEEE Spectrum piece (2 360 w).",
         "IBM Research podcast (CC).",
         "Horizon scan (520 w)",
         "Side channels en nuevos c\u00f3mputos"),
        ("Roundtable: emerging tech",
         "Polished discussion",
         "+50: let me build on, I would like to pick up on, with due respect, I would nuance that, push back gently, lean into, pivot to, zoom out, zoom in, operationalize, strawman, steelman, intellectual honesty, epistemic humility, strong priors, update, Bayesian",
         "Transcript OECD AI roundtable (2 370 w).",
         "OECD Future of Work podcast (CC).",
         "Opening + closing statement (530 w)",
         "Negotiating multilateral frames"),
        ("Mock CPE + master capstone",
         "Integrado",
         "16 000+ acumuladas",
         "CPE Reading & UoE completo (~5 500 w, 7 parts).",
         "CPE Listening completo (40 min).",
         "Essay (240\u2013280 w) + review/report/proposal (280\u2013320 w)",
         "Entregable: whitepaper publicable"),
    ]),
]

# -------- render de cada nivel --------
for code, desc, quota, classes in LEVELS:
    add_heading(f"5.{LEVELS.index((code, desc, quota, classes)) + 1}  Nivel {code} \u2013 {desc}", 2)
    add_para(f"Objetivo l\u00e9xico: {quota}.", italic=True)

    # tabla de 20 filas con columnas
    headers = ["#", "Tema", "Gram\u00e1tica / Use of English", "Vocabulario nuevo (muestra)",
               "Reading (Cambridge / TOEFL-IELTS)", "Listening (fuente libre)",
               "Writing (palabras)", "Ciber / Tech angle"]

    rows = []
    for i, (tema, gram, vocab, reading, listening, writing, angle) in enumerate(classes, 1):
        rows.append([str(i), tema, gram, vocab, reading, listening, writing, angle])

    t = add_table(headers, rows,
                  col_widths=[0.6, 2.5, 2.3, 3.5, 3.2, 2.5, 1.6, 1.8])
    # fuente m\u00e1s chica para tablas densas
    for row in t.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(8)
    page_break()

# =========================================================
# 6. EVALUACION Y CERTIFICACION
# =========================================================
add_heading("6. Evaluaci\u00f3n y certificaci\u00f3n interna", 1)
add_para(
    "Cada nivel cierra con un examen sumativo en formato Cambridge completo. La calificaci\u00f3n "
    "interna se reporta como Cambridge Scale (100\u2013230) con nota m\u00ednima de promoci\u00f3n de 160 "
    "para niveles A1\u2013B1 y 180 para B2\u2013C2. Se emite constancia interna con mapeo CEFR, "
    "Cambridge, IELTS y TOEFL iBT."
)
add_table(
    ["Componente", "Peso", "Formato"],
    [
        ["Reading & Use of English", "25%", "Cambridge + 1 texto TOEFL-style"],
        ["Listening", "25%", "Cambridge + audios del nivel"],
        ["Writing", "25%", "2 tareas seg\u00fan banda CEFR"],
        ["Speaking", "15%", "Examen oral en pares + interview"],
        ["Portafolio / Capstone", "10%", "Producto final del nivel"],
    ],
    col_widths=[5.0, 2.0, 10.0],
)

doc.add_paragraph()
add_heading("6.1 R\u00fabricas", 2)
add_para(
    "Se usan las r\u00fabricas oficiales de Cambridge (Content, Communicative Achievement, "
    "Organisation, Language) para Writing y Speaking, adaptadas al dominio de sistemas. "
    "Para Reading se usa item analysis (facility / discrimination) y para Listening se "
    "registra velocidad y acento del audio para asegurar variabilidad."
)

# =========================================================
# 7. BIBLIOGRAFIA Y RECURSOS
# =========================================================
add_heading("7. Bibliograf\u00eda y recursos de referencia", 1)
add_para("Marcos y cat\u00e1logos l\u00e9xicos", bold=True)
for x in [
    "Council of Europe (2020). CEFR Companion Volume with New Descriptors.",
    "English Profile (Cambridge) \u2013 English Vocabulary Profile Online (EVP).",
    "Cambridge Assessment English \u2013 Handbooks for Teachers A1 Key, B1 Preliminary, B2 First, C1 Advanced, C2 Proficiency (ediciones vigentes).",
    "ETS \u2013 TOEFL iBT Test Prep Planner.",
    "IELTS Partnership \u2013 Academic Reading / Listening guides.",
]: doc.add_paragraph(x, style='List Bullet')

add_para("Dominio t\u00e9cnico (lecturas de \u00faltima tecnolog\u00eda)", bold=True)
for x in [
    "NIST SP 800-207 (Zero Trust), 800-61r2 (IR), 800-53r5 (Controls), 800-63B (Identity).",
    "CISA Known Exploited Vulnerabilities catalog + Cybersecurity Advisories.",
    "ENISA Threat Landscape (\u00faltima edici\u00f3n).",
    "OWASP Top 10 (Web 2021, LLM 2025, API 2023, Mobile 2024).",
    "IEEE Spectrum, IEEE Security & Privacy, ACM Queue, Communications of the ACM.",
    "arXiv.org / cs.CR (cryptography & security).",
    "MIT Technology Review, The Register, Krebs on Security, SANS ISC Diary.",
]: doc.add_paragraph(x, style='List Bullet')

add_para("Audio libre (revise siempre la licencia del recurso espec\u00edfico antes de distribuirlo)", bold=True)
for x in [
    "VOA Learning English (learningenglish.voanews.com) \u2013 Dominio p\u00fablico (US gov).",
    "NASA audio (nasa.gov/audio) \u2013 Dominio p\u00fablico.",
    "CISA, NIST, FBI podcasts \u2013 Dominio p\u00fablico.",
    "LibriVox (librivox.org) \u2013 Dominio p\u00fablico.",
    "Internet Archive (archive.org) \u2013 filtrar por licencia CC / PD.",
    "MIT OpenCourseWare (ocw.mit.edu) \u2013 CC BY-NC-SA.",
    "OpenLearn (open.edu/openlearn) \u2013 CC BY-NC-SA.",
    "EFF Podcasts (eff.org) \u2013 CC BY.",
    "Freesound.org \u2013 filtrar por CC0 / CC BY.",
    "Pixabay audio \u2013 Pixabay License (libre con cond. propias).",
]: doc.add_paragraph(x, style='List Bullet')

# =========================================================
# 8. APENDICE: PALABRAS POR NIVEL (criterios)
# =========================================================
page_break()
add_heading("8. Ap\u00e9ndice \u2013 criterios de selecci\u00f3n l\u00e9xica", 1)
add_para(
    "Las listas de vocabulario muestran *ejemplos representativos*; la lista completa por nivel "
    "se entrega en hoja Excel aparte para facilitar el trabajo con spaced repetition (Anki, "
    "Quizlet, Memrise). Se combinan:"
)
for x in [
    "General Service List (GSL-NG) + NGSL de Browne, Culligan & Phillips (2013+).",
    "Academic Word List (Coxhead) + New Academic Vocabulary List (Gardner & Davies).",
    "English Vocabulary Profile (Cambridge) por banda A1\u2013C2.",
    "Listas especializadas: Coxhead Engineering Word List, NIST Cybersecurity Glossary (NISTIR 7298r3), ENISA Glossary.",
    "Frecuencia en corpus t\u00e9cnico actual (arXiv cs.CR, Stack Overflow) para cubrir jerga viva.",
]: doc.add_paragraph(x, style='List Bullet')

# =========================================================
# CIERRE
# =========================================================
add_heading("9. Cierre", 1)
add_para(
    "Este programa es la base para que el Departamento de Ingl\u00e9s de Nordic International "
    "School (NIS Per\u00fa) o cualquier escuela t\u00e9cnica o Facultad de Ingenier\u00eda pueda desplegar "
    "un itinerario de ingl\u00e9s t\u00e9cnico A1\u2192C2 alineado al CEFR y a los ex\u00e1menes "
    "internacionales Cambridge, IELTS y TOEFL. Los 120 planes de clase pueden expandirse "
    "a cuadernos de estudiante y gu\u00eda docente con 6\u20138 p\u00e1ginas por clase."
)

doc.save(OUTFILE)
print("OK", OUTFILE)
